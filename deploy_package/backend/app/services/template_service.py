"""
文书模板业务逻辑
"""
from typing import List, Optional, Dict, Any
from motor.motor_asyncio import AsyncIOMotorDatabase
import re
import math
from datetime import datetime
from jinja2 import Template
from io import BytesIO

# 可选依赖：PDF 和 Word 导出功能
try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class TemplateService:
    """文书模板服务类"""

    def __init__(self, db: AsyncIOMotorDatabase):
        self.db = db
        self.templates_collection = db.doc_templates
        self.instances_collection = db.doc_instances

    async def get_templates_list(
        self,
        page: int = 1,
        page_size: int = 20,
        category: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        获取模板列表
        """
        query = {}
        if category:
            query["category"] = category

        total = await self.templates_collection.count_documents(query)

        skip = (page - 1) * page_size
        total_pages = math.ceil(total / page_size) if total > 0 else 0

        cursor = self.templates_collection.find(query).skip(skip).limit(page_size)
        templates = await cursor.to_list(length=page_size)

        for template in templates:
            template["_id"] = str(template["_id"])

        return {
            "data": templates,
            "pagination": {
                "total": total,
                "page": page,
                "page_size": page_size,
                "total_pages": total_pages,
            }
        }

    async def get_template_detail(self, template_id: str) -> Optional[Dict[str, Any]]:
        """
        获取模板详情
        """
        template = await self.templates_collection.find_one({"template_id": template_id})
        if template:
            template["_id"] = str(template["_id"])
        return template

    async def create_template(self, template_data: Dict[str, Any]) -> str:
        """
        创建模板
        """
        template_data["created_at"] = datetime.utcnow()
        result = await self.templates_collection.insert_one(template_data)
        return str(result.inserted_id)

    async def update_template(self, template_id: str, template_data: Dict[str, Any]) -> bool:
        """
        更新模板
        """
        result = await self.templates_collection.update_one(
            {"template_id": template_id},
            {"$set": template_data}
        )
        return result.modified_count > 0

    async def delete_template(self, template_id: str) -> bool:
        """
        删除模板
        """
        result = await self.templates_collection.delete_one({"template_id": template_id})
        return result.deleted_count > 0

    async def render_template(self, template_id: str, field_values: Dict[str, Any]) -> str:
        """
        渲染模板（填充变量）
        """
        template = await self.get_template_detail(template_id)
        if not template:
            raise ValueError(f"模板不存在: {template_id}")

        # 使用 Jinja2 模板引擎
        jinja_template = Template(template["content"])
        rendered_content = jinja_template.render(**field_values)

        return rendered_content

    async def export_to_pdf(
        self, template_id: str, field_values: Dict[str, Any]
    ) -> bytes:
        """
        导出为 PDF
        """
        if not WEASYPRINT_AVAILABLE:
            raise RuntimeError("PDF 导出功能不可用：WeasyPrint 未安装")
        
        # 渲染模板
        rendered_html = await self.render_template(template_id, field_values)

        # 添加 CSS 样式
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: "Noto Sans CJK SC", "SimSun", serif;
                    font-size: 14pt;
                    line-height: 1.8;
                }}
                h1 {{
                    text-align: center;
                    font-size: 20pt;
                    margin-bottom: 1cm;
                }}
                .field {{
                    margin: 0.5cm 0;
                }}
                .signature {{
                    text-align: right;
                    margin-top: 2cm;
                }}
            </style>
        </head>
        <body>
            {rendered_html}
        </body>
        </html>
        """

        # 生成 PDF
        pdf_bytes = HTML(string=styled_html).write_pdf()

        # 保存实例记录
        await self._save_instance(template_id, field_values, "pdf")

        return pdf_bytes

    async def export_to_docx(
        self, template_id: str, field_values: Dict[str, Any]
    ) -> bytes:
        """
        导出为 DOCX
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("Word 导出功能不可用：python-docx 未安装")
        
        # 渲染模板
        rendered_content = await self.render_template(template_id, field_values)

        # 创建 Word 文档
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(14)

        # 获取模板信息
        template = await self.get_template_detail(template_id)

        # 添加标题
        title = doc.add_paragraph(template["name"])
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.runs[0].font.size = Pt(20)
        title.runs[0].font.bold = True

        # 添加内容（按行分割）
        for line in rendered_content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)

        # 保存到内存
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)

        # 保存实例记录
        await self._save_instance(template_id, field_values, "docx")

        return docx_io.getvalue()

    async def _save_instance(
        self, template_id: str, field_values: Dict[str, Any], export_format: str
    ):
        """
        保存文书实例记录
        """
        instance_data = {
            "template_id": template_id,
            "field_values": field_values,
            "export_format": export_format,
            "created_at": datetime.utcnow()
        }
        await self.instances_collection.insert_one(instance_data)

    async def get_categories(self) -> List[str]:
        """
        获取所有模板分类
        """
        categories = await self.templates_collection.distinct("category")
        return sorted(categories)
